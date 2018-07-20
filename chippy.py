"""
Chippy Dictation
ver: 0.9 
by: Irza Pulungan

requirement

pythoncom
  conda install -c anaconda comtypes 

pyHoook
  download ke
    https://www.lfd.uci.edu/~gohlke/pythonlibs/
  ambil:
    python_hdf4‑0.9‑cp27‑cp27m‑win_amd64.whl
  install:
    pip.exe install pyHook-1.5.1-cp27-cp27m-win_amd64.whl

selenium
  conda install -c conda-forge selenium 

python-docx  
  conda install -c conda-forge python-docx 

chromedriver:
    https://sites.google.com/a/chromium.org/chromedriver/

pyinstaller(optional)    
   conda install -c conda-forge pyinstaller   
 

    
"""
import pythoncom, pyHook
import ctypes
import re
import datetime
import os
import time
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import sys
import signal
from docx import Document
import win32event, win32api, winerror
from time import gmtime, strftime
import yaml


#Disallowing Multiple Instance
mutex = win32event.CreateMutex(None, 1, 'mutex_var_chippy')
if win32api.GetLastError() == winerror.ERROR_ALREADY_EXISTS:
    mutex = None
    print "Multiple Instance not Allowed"
    exit(0)
x=''
data=''
count=0

#ALL SETTINGS
#set debug
debug = False
apppath = os.path.dirname(os.path.realpath(sys.argv[0]))
ymlfile = os.path.join(apppath,'config.yml')

try:
    appconfig = yaml.load(open(ymlfile,'r'))
    docdir = appconfig['docdir']
except:
    docdir = apppath
    pass

filename = strftime("%Y-%m-%d-%H%M%S-vlog", gmtime()) + '.docx'
#filename = '\dictation.docx'
filename = os.path.join(docdir,filename)
talk_state = False
sleep_state = False




def OnKeyboardEvent(event):
    global browser, debug
    if debug:
      print 'MessageName:',event.MessageName
      print 'Message:',event.Message
      print 'Time:',event.Time
      print 'Window:',event.Window
      print 'WindowName:',event.WindowName
      print 'Ascii:', event.Ascii, chr(event.Ascii)
      print 'Key:', event.Key
      print 'KeyID:', event.KeyID
      print 'ScanCode:', event.ScanCode
      print 'Extended:', event.Extended
      print 'Injected:', event.Injected
      print 'Alt', event.Alt
      print 'Transition', event.Transition
      print '---'

    if event.Key == 'F10':
       browser.quit()
       talk("dadah")
       MessageBox = ctypes.windll.user32.MessageBoxA
       MessageBox(None, 'The Chippy Program exited..., file disimpan di '+filename, 'alert', 0)
       print "Exiting..."
       sys.exit()
  
    if event.Key == 'F9':
       get_dict() 

    if event.Key == 'Oem_Period':
       a = get_content()
       a = a + "."
       save_content(a)
       time.sleep(0.1)
       get_dict()
        
    
    if event.Key == 'Oem_Comma':
       a = get_content()
       a = a + ","
       save_content(a)
       time.sleep(0.1)
       get_dict()

    
    if event.Key == 'Oem_2':
        pass
    
    
    
        

# return True to pass the event to other handlers
    return True


def get_dict():
      source = browser.find_element_by_xpath('//*[@id="source"]')
      source.clear()
      record = browser.find_element_by_xpath('//*[@id="gt-speech"]/span')
      record.click()
      #time.sleep(5)
      #record
      #record.click()
      #a = source.get_attribute("value")
      #print a
      #listen_state = False
      #source.clear()
      return


def get_content():
    record = browser.find_element_by_xpath('//*[@id="gt-speech"]/span')
    record.click()
    source = browser.find_element_by_xpath('//*[@id="source"]')
    a = source.get_attribute("value")
    return a
    
def save_content(a):
    global document
    paragraph = document.add_paragraph(a)
    document.save(filename)
    
def talk(speech):
    global talk_state, listen_state, sleep_state ,yesno_state, cmd, prm
    if not talk_state:
        if not sleep_state:
            try:
                talk_state = True
                tts = browser.find_element_by_xpath('//*[@id="source"]')
                tts.clear()
                tts.send_keys(str(speech))
                time.sleep(1)
                talk = browser.find_element_by_xpath('//*[@id="gt-src-listen"]/span')
                talk.click()
                time.sleep(2)
                talk_state = False
            except:
                pass



chrome_options = Options()
chrome_options.add_argument("--use-fake-ui-for-media-stream")
chrome_options.add_argument('--ignore-certificate-errors')
chrome_options.add_argument("--test-type")

browser = webdriver.Chrome(chrome_options=chrome_options)
browser.get('https://translate.google.com/?#id/en/Tekan%20F9%20untuk%20rekam,F10%20untuk%20QUIT')
assert "Google Translate" in browser.title

try:
  document = Document(filename)
except:
  document = Document()


# create a hook manager
hm = pyHook.HookManager()
# watch for all mouse events
hm.KeyDown = OnKeyboardEvent
# set the hook
hm.HookKeyboard()
# wait forever
pythoncom.PumpMessages()


